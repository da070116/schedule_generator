import docx

from docx.shared import RGBColor
from PyQt5 import QtWidgets
from converterui import Ui_MainWindow


def clarify(acc_string):
    return acc_string.replace('́', '').replace('\n', '<br>').replace('\xa0', '').replace('.00', ':00')


def docx_to_html(doc: docx.Document):
    html_content = ''
    for counter, paragraph in enumerate(doc.paragraphs):
        html_content = html_content + f"<p>{paragraph.text}</p>\n"
    table = doc.tables[0]
    html_content = html_content + "<table>\n"
    for counter_rows, row in enumerate(table.rows):
        if counter_rows > 0:
            html_content = html_content + "<tr>\n"
            for cell in row.cells:
                color_rgb = None
                is_bold = None
                if cell.paragraphs:
                    color_rgb = cell.paragraphs[0].runs[0].font.color.rgb
                    if cell.paragraphs[0].runs[0].font.bold:
                        is_bold = True
                fixed_text = clarify(cell.text)
                if color_rgb == docx.shared.RGBColor(0xff, 0x00, 0x00):
                    fixed_text = f'<span style="color:#ff0000;">{fixed_text}</span>'
                if is_bold:
                    fixed_text = f'<strong>{fixed_text}</strong>'
                fixed_text = f'<td>{fixed_text}</td>\n'
                html_content = html_content + fixed_text
            html_content = html_content + "</tr>\n"
    html_content = html_content + "</table>\n"
    return html_content


class ConverterApp(QtWidgets.QMainWindow, Ui_MainWindow):
    def __init__(self):
        super().__init__()
        self.setupUi(self)
        self.schedule_file = ""

        self.pickButton.pressed.connect(self.pick_document)
        self.convertButton.pressed.connect(self.convert)

    def pick_document(self):
        self.schedule_file = QtWidgets.QFileDialog.getOpenFileName(self, "Pick file")[0]
        if self.schedule_file:
            self.lineEdit.setText(self.schedule_file)
            self.convertButton.setDisabled(False)


    def convert(self):
        if self.schedule_file:
            document_to_parse = docx.Document(self.schedule_file)
            content_file = docx_to_html(document_to_parse)
            with open('ready.html', 'wt', encoding='utf-8') as f:
                f.write(content_file)
            self.lineEdit.setText("")
            self.convertButton.setDisabled(True)


app = QtWidgets.QApplication([])
window = ConverterApp()
window.show()
app.exec_()
